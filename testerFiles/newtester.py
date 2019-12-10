import docx
from openpyxl import load_workbook
from docx.shared import Pt, RGBColor
import os

# import shutil

# dir_name = "katilimcilar"
complete_path = "C:\\Users\\ASUS\\PycharmProjects\\pdfWriter\\katilimcilar"
# if os.path.exists(os.path.dirname(dir_name)):
#     print("exist")
#     shutil.rmtree(complete_path)
os.chdir(complete_path)

workbook = load_workbook(filename="tester\\temp.xlsx")
sheet = workbook.active
names = []
surnames = []
schools = []
counter = 0
# reading excel file.
for value in sheet.iter_cols(min_row=1, max_col=2, values_only=True):
    for i in range(1, len(value)):
        if counter == 0:
            names.append(value[i])
        # elif counter == 1:
        #     surnames.append(value[i])
        elif counter == 1:
            schools.append(value[i])
    counter += 1

document = docx.Document("tester\\beyazfon.docx")
counter = 0
# ======================== DEBUG STARTS ========================


# 11 Name runs[0] => isim, runs[2] => soyisim
# 14 üniversite adı runs[0] => Sabanci Universitesi

nameParagraphIdx = 11
schoolParagraphIdx = 14

# counter = 0
# for i in document.paragraphs[nameParagraphIdx].runs:
#     print(counter, i.text)
#     counter += 1

print(document.paragraphs[schoolParagraphIdx].runs[0].text)

# ======================== DEBUG ENDS ========================

para_format = document.paragraphs[nameParagraphIdx].paragraph_format
nameRun = document.paragraphs[nameParagraphIdx].runs[0]
# nameRun = document.paragraphs[nameParagraphIdx].runs[0]
print("\n******* Name Font *******\n")
nameFont = nameRun.font
print("Before: ", nameFont.color.rgb)
nameFont.color.rgb = RGBColor(0x00, 0x27, 0x76)
print("After: ", nameFont.color.rgb)
nameFont.size = Pt(18)
# fontName = nameRun.font
# fontName.color.rgb = RGBColor(0x00, 0x27, 0x76)
# fontName.size = Pt(18)

# para_format2 = document.paragraphs[nameParagraphIdx].paragraph_format
# surnameRun = document.paragraphs[nameParagraphIdx].runs[2]
# fontSurname = surnameRun.font
# fontSurname.color.rgb = RGBColor(0x00, 0x27, 0x76);
# fontSurname.size = Pt(18)

# surname_run = document.paragraphs[6].runs[0]
# font_surname = surname_run.font
# font_surname.color.rgb = RGBColor(0x00, 0x27, 0x76);
# font_surname.size = Pt(18)

# school_run = document.paragraphs[schoolParagraphIdx].runs[0]
print("\n******* School Font *******")
school_run = document.paragraphs[schoolParagraphIdx].runs[0]
font_school = school_run.font
font_school.color.rgb = RGBColor(0x00, 0x27, 0x76);
font_school.size = Pt(15)


def buyukHarfCevir(sStr):

    aranan = ''
    HARFDIZI = [
        ('i', 'İ'), ('ğ', 'Ğ'), ('ü', 'Ü'), ('ş', 'Ş'), ('ö', 'Ö'), ('ç', 'Ç'),
        ('ı', 'I')
    ]
    for aranan, harf in HARFDIZI:
        sStr = sStr.replace(aranan, harf)

    sStr = sStr[0].capitalize() + sStr[1:].lower()
    return sStr


curr_student = 0
while curr_student < len(names) - 1:

    # name of the current student
    name = names[curr_student]
    if name != None:
        name = name.replace(" ", "-")

    # school of the current student
    school = schools[curr_student]
    if school != None:
        school = school.replace(" ", "-")

    # split current name in order to get surname and name separately
    curNameList = names[curr_student].split()

    # last element of the list is the surname of current attendee.
    surname = curNameList[-1]

    # buraya ilk harf türkçe karakterse kontrolu yap. !!!!!!!!!!1
    surname = surname.capitalize()

    listLen = len(curNameList);
    curNameList.pop(listLen - 1)

    curName = curNameList[0]
    name = name.lower()
    curName = curName.capitalize()

    # ==== TESTER NAME and SCHOOL TEXT STARTS ====
    if len(surname) > 12:
        print("Length of the surname is greater than 12:", surname)
        nameRun.font.size = Pt(17)

    nameRun.font.size.size = Pt(18)
    document.paragraphs[nameParagraphIdx].runs[0].text = curName + " " + surname
    document.paragraphs[schoolParagraphIdx].runs[0].text = schools[curr_student]

    # ==== TESTER NAME and SCHOOL TEXT ENDS ====
    # document.paragraphs[schoolParagraphIdx].runs[2].text = ""
    # document.paragraphs[schoolParagraphIdx].runs[0].text = schools[curr_student].capitalize

    documentName = name + "-" + school + ".docx"
    curr_student += 1
    document.save(documentName)
