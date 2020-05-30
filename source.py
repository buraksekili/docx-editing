import docx
from openpyxl import load_workbook
from docx.shared import Pt, RGBColor

workbook = load_workbook(filename="temp.xlsx")
sheet = workbook.active
names = []
surnames = []
schools = []
counter = 0

for value in sheet.iter_cols(min_row=1, max_col=2, values_only=True):
    for i in range(1, len(value)):
        if counter == 0:
            names.append(value[i])
        elif counter == 1:
            schools.append(value[i])
    counter += 1

document = docx.Document("tester1.docx")

para_format = document.paragraphs[5].paragraph_format
nameRun = document.paragraphs[5].runs[0]
fontName = nameRun.font
fontName.color.rgb = RGBColor(0x00, 0x27, 0x76)
fontName.size = Pt(18)

para_format2 = document.paragraphs[6].paragraph_format
surnameRun = document.paragraphs[6].runs[0]
fontSurname = surnameRun.font
fontSurname.color.rgb = RGBColor(0x00, 0x27, 0x76)
fontSurname.size = Pt(18)

school_run = document.paragraphs[7].runs[0]
font_school = school_run.font
font_school.color.rgb = RGBColor(0x00, 0x27, 0x76)
font_school.size = Pt(12)


def buyukHarfCevir(sStr):
    str_copy = sStr
    HARFDIZI = [
        ('i', 'İ'), ('ğ', 'Ğ'), ('ü', 'Ü'), ('ş', 'Ş'), ('ö', 'Ö'), ('ç', 'Ç'),
        ('ı', 'I')
    ]
    for aranan, harf in HARFDIZI:
        str_copy = str_copy.replace(aranan, harf)

    str_copy = str_copy[0] + str_copy[1:].lower()
    return str_copy


curr_student = 0
while curr_student < len(names) - 1:
    name = names[curr_student]
    if name != None:
        name = name.replace(" ", "-")

    school = schools[curr_student]
    if school != None:
        school = school.replace(" ", "-")
    curNameList = names[curr_student].split()

    surname = curNameList[-1]
    surname = surname.capitalize()

    listLen = len(curNameList)
    curNameList.pop(listLen - 1)

    curName = ""
    for i in curNameList:
        curName += buyukHarfCevir(i) + " "
    document.paragraphs[5].runs[0].text = curName
    if len(surname) > 12:
        print(surname)
        fontSurname.size = Pt(17)
        print(fontSurname.size)
    document.paragraphs[6].runs[0].text = surname
    fontSurname.size = Pt(18)
    if schools[curr_student].count(' ') == 0:
        document.paragraphs[7].runs[2].text = ""
        document.paragraphs[7].runs[0].text = schools[curr_student]
    else:
        document.paragraphs[7].runs[2].text = ""
        document.paragraphs[7].runs[0].text = schools[curr_student]
    documentName = name + "-" + school + ".docx"
    curr_student += 1
    document.save(documentName)
    
